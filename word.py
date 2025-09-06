from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Set up the title
title = doc.add_heading('Mind Map', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add placeholder for central photo
photo_paragraph = doc.add_paragraph()
photo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = photo_paragraph.add_run()
run.add_picture("", width=Inches(1.5))  # Replace with local placeholder if needed

photo_caption = doc.add_paragraph("<< Your Photo Here >>")
photo_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
photo_caption.runs[0].italic = True

# Define branches (main categories)
branches = ["Branch 1", "Branch 2", "Branch 3", "Branch 4"]
branch_colors = ["FF9999", "99CCFF", "99FF99", "FFCC99"]

for idx, branch in enumerate(branches):
    branch_title = doc.add_paragraph()
    run = branch_title.add_run(branch)
    run.bold = True
    run.font.size = Pt(14)
    branch_title.paragraph_format.left_indent = Inches(0.25 * (idx + 1))
    run.font.color.rgb = RGBColor.from_string(branch_colors[idx])

    # Add leaves for each branch
    for i in range(3):
        leaf = doc.add_paragraph(f"<< Leaf {i+1} for {branch} >>")
        leaf.paragraph_format.left_indent = Inches(0.4 * (idx + 1))
        leaf.style = 'List Bullet'

# Save document
doc.save("Mind_Map_Template.docx")
print("✅ Word file created: Mind_Map_Template.docx")
